from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent


DOCS = [
    (ROOT / "docs" / "技术规范文档.md", ROOT / "技术规范文档.docx"),
    (ROOT / "docs" / "可行性研究报告.md", ROOT / "可行性研究报告.docx"),
]


def set_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def shade(cell, color: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def clean_inline(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("__", "")
        .replace("`", "")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .replace("<br />", "\n")
    )


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    chars = set(stripped.replace("|", "").replace(":", "").replace("-", "").replace(" ", ""))
    return not chars and "-" in stripped


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [clean_inline(cell.strip()) for cell in stripped.split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx in range(width):
        table.rows[0].cells[idx].text = rows[0][idx] if idx < len(rows[0]) else ""
        shade(table.rows[0].cells[idx])

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(width):
            cells[idx].text = row[idx] if idx < len(row) else ""
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def add_paragraph(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    if stripped == "---":
        doc.add_paragraph("—" * 24)
        return

    if stripped.startswith("#"):
        level = len(stripped) - len(stripped.lstrip("#"))
        text = clean_inline(stripped[level:].strip())
        if level == 1:
            doc.add_heading(text, level=0)
        else:
            doc.add_heading(text, level=min(level - 1, 4))
        return

    if stripped.startswith(">"):
        p = doc.add_paragraph(clean_inline(stripped.lstrip(">").strip()))
        p.style = "Intense Quote" if "Intense Quote" in doc.styles else doc.styles["Normal"]
        return

    if stripped.startswith(("- ", "* ")):
        doc.add_paragraph(clean_inline(stripped[2:].strip()), style="List Bullet")
        return

    if len(stripped) > 3 and stripped[0].isdigit():
        dot = stripped.find(".")
        if 0 < dot <= 3 and stripped[dot + 1 : dot + 2] == " ":
            doc.add_paragraph(clean_inline(stripped[dot + 2 :].strip()), style="List Number")
            return

    doc.add_paragraph(clean_inline(stripped))


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    set_style(doc)
    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines):
                row_line = lines[i]
                row_stripped = row_line.strip()
                if not (row_stripped.startswith("|") and row_stripped.endswith("|")):
                    break
                if not is_table_separator(row_stripped):
                    table_rows.append(split_table_row(row_stripped))
                i += 1
            add_table(doc, table_rows)
            continue

        add_paragraph(doc, line)
        i += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(docx_path)


def main() -> None:
    for md_path, docx_path in DOCS:
        if not md_path.exists():
            raise FileNotFoundError(md_path)
        convert(md_path, docx_path)
        print(f"generated: {docx_path.name}")


if __name__ == "__main__":
    main()

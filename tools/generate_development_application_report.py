from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "开发与应用报告.docx"
TITLE = "开发与应用报告"
PROJECT = "基于特色创新项目的岭南红橙病虫害智能监测与校本教研一体化平台"


def set_font(run, font: str, size: int, bold: bool = False) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def para(p, first_line: bool = False) -> None:
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.85)


def add_text(doc: Document, text: str, font="仿宋_GB2312", size=16, bold=False, center=False, first_line=False) -> None:
    p = doc.add_paragraph()
    para(p, first_line=first_line)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, font, size, bold)


def add_title(doc: Document) -> None:
    add_text(doc, TITLE, "方正小标宋简体", 18, center=True)


def add_h1(doc: Document, text: str) -> None:
    add_text(doc, text, "黑体", 16)


def add_h2(doc: Document, text: str) -> None:
    add_text(doc, text, "楷体_GB2312", 16)


def add_body(doc: Document, text: str) -> None:
    add_text(doc, text, "仿宋_GB2312", 16, first_line=True)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        add_text(doc, line, "仿宋_GB2312", 16)


def configure(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(16)


def generate() -> None:
    doc = Document()
    configure(doc)
    add_title(doc)

    add_h1(doc, "一、开发背景")
    add_body(
        doc,
        "本案例面向高职涉农物联网专业实训教学中的真实问题展开。传统课堂中，病虫害识别多依赖教材图片和教师经验讲解，"
        "学生难以接触真实复杂农情；科研项目中的一线数据和算法成果转化为教学资源的周期较长；学生实训过程、识别能力、"
        "防治方案设计和台账整理等表现也缺少多维评价依据。为破解这些痛点，项目依托广东省特色创新类项目"
        "（2025KQNCX312），将岭南红橙种植生产一线的病虫害图像、物联网温湿度等环境数据和校本防治知识库整合起来，"
        "开发“智能监测与校本教研一体化平台”，推动真实生产数据进入课堂。"
    )
    add_body(
        doc,
        "平台的目标不是单纯完成图像识别，而是形成“真实图片导入—本地AI识别—环境物候补充—知识库推理—台账沉淀—"
        "报表导出—教学案例生成”的闭环，使教师能够快速把科研成果转化为案例教学、农技实训和过程评价资源。"
    )

    add_h1(doc, "二、设计与开发")
    add_h2(doc, "（一）平台与技术选择")
    add_body(
        doc,
        "系统采用 Python 3.11、PyQt5、OpenCV、SQLite3、openpyxl、reportlab、ONNX Runtime 和 YOLOv8s-XH 等技术。"
        "前端使用 PyQt5 构建大字体、高对比度、单键触发的适老化桌面界面；核心算法采用本地部署的 YOLOv8s-XH 模型，"
        "实现红橙黄龙病、红蜘蛛、溃疡病等12类高发病虫害的离线识别与群聚计数；数据层使用 SQLite3 构建校本知识库、"
        "历史台账和农户档案；报表层支持 Excel、PDF 和 Markdown 离线导出；教研层支持模板生成和可选大模型增强。"
    )

    add_h2(doc, "（二）借助生成式人工智能的开发过程")
    add_body(
        doc,
        "开发过程中，生成式人工智能主要用于需求重构、模块设计、代码生成、文档生成和提示词优化。首先，将“涉农物联网专业"
        "实训教学痛点”“省特色创新项目成果转化”“红橙病虫害识别”等自然语言需求输入 AI，辅助拆解为功能模块：智能检测、"
        "物候期选择、知识库推理、检测台账、报表导出、教学生成和设置诊断。其次，借助 AI 生成并迭代 PyQt 页面、SQLite "
        "表结构、YOLO 推理封装、Excel/PDF 导出脚本和 docx/pptx 文档生成脚本，再由开发者结合项目运行结果进行调试修正。"
    )
    add_body(
        doc,
        "生成式 AI 参与开发的证据主要体现在三类可复现产物中：一是项目中保留了多个 Python 自动化脚本，如"
        "generate_project_docx.py、convert_markdown_docs_to_docx.py、generate_development_application_report.py、"
        "generate_demo_pptx_from_template.py，可直接生成安装手册、使用手册、开发报告和演示PPT；二是教研生成模块中设计了"
        "面向大模型的 Prompt 和降级模板，把检测结果、物候期、知识库处方转化为教学案例；三是系统设置页提供 OpenAI 兼容"
        "接口，可连接本地大模型或 DeepSeek、豆包等云端模型。"
    )
    add_code(
        doc,
        "提示词证据示例：\n"
        "你是岭南红橙现代农业产业园的农技教研专家。请根据本地病虫害检测结果与已审核绿色防治知识库处方，"
        "生成结构清晰的中文教学文档；必须突出化学药剂安全间隔期PHI；若主病害为黄龙病，必须明确确诊植株立即砍除并就地烧毁。"
    )
    add_body(
        doc,
        "通过上述方式，生成式人工智能不是替代开发者，而是作为“需求分析助手、代码草稿助手、文档生成助手和教研内容助手”。"
        "所有 AI 生成内容均经过人工审核，涉及病虫害防治、药剂安全间隔期和教学评价的数据均以本地知识库和教师审核为准。"
    )

    add_h2(doc, "（三）功能架构")
    add_body(
        doc,
        "系统最终形成七类功能：一是图片、视频、摄像头病虫害智能识别；二是基于物候期和环境数据的绿色防治知识库推理；"
        "三是农户档案、检测日志和历史台账管理；四是 Excel、PDF、Markdown 报表导出；五是数据采集、标注、训练和训练结果"
        "可视化；六是教学案例和学生农技实训指导意见生成；七是模型、字体、阈值、保存目录和大模型接口设置。核心闭环为："
        "导入真实病虫害图片，YOLOv8s-XH 本地识别，补充物候期和温湿度，知识库生成防治方案，保存检测台账，导出报表，"
        "生成教学案例。"
    )

    add_h1(doc, "三、应用过程与效果")
    add_body(
        doc,
        "在课堂演示中，教师首先导入一张来自岭南红橙生产一线的病虫害图片，系统在本地完成识别，显示检测框、类别、置信度、"
        "严重程度和群聚计数。随后选择红橙物候期，并可接入温湿度等物联网环境数据。系统根据“病虫害类别+物候期+环境数据”"
        "匹配 SQLite 校本知识库，生成物理、生物、合规化学三位一体防治方案，并突出安全间隔期 PHI。最后，检测过程自动保存"
        "为农情台账，教师可导出 Excel/PDF 报表，也可生成《岭南红橙实时防害教学案例》或《学生农技实训指导意见》。"
    )
    add_body(
        doc,
        "应用成效主要体现在四个方面。第一，提升备课效率：以往教师需要手工查找图片、整理症状、编写实训任务和防治方案，"
        "现在可由一次检测自动生成案例草稿和台账报表。第二，增强学习真实感：学生面对的是生产一线图像和环境数据，而不是"
        "静态教材样图，能够在真实复杂情境中训练观察、识别、判断和记录能力。第三，支撑过程评价：系统沉淀检测时间、图像、"
        "病虫害类别、物候期、处方建议和导出记录，可作为学生实训过程和成果评价依据。第四，扩大教研影响力：平台把省特色"
        "创新项目成果转化为可演示、可复制、可归档的数字教研资产，可服务课堂教学、实训室演示、技能竞赛训练和农技推广。"
    )
    add_body(
        doc,
        "从使用影响看，平台降低了高职涉农物联网专业开展 AI+农业实训的技术门槛。系统默认离线运行，普通办公电脑和 USB "
        "摄像头即可演示；Excel 报表可快速离线导出，便于教研归档；教学生成模块让教师能够把一次真实检测转化为课堂案例，"
        "形成“科研成果—教学资源—学生实训—评价数据”的持续循环。"
    )

    add_h1(doc, "四、创新与反思")
    add_body(
        doc,
        "本成果的创新点在于科教融合、边缘智能和教研闭环。科教融合方面，项目把省特色创新项目中的真实农情数据转化为"
        "校本数字教研资产；边缘智能方面，构建“YOLOv8s-XH 本地识别+SQLite 知识库推理”的离线架构，降低网络依赖和运行成本；"
        "教研闭环方面，把检测结果进一步转化为教学案例、实训指导和评价台账，突破了传统识图工具只识别不教学的问题。"
    )
    add_body(
        doc,
        "反思来看，平台仍需持续完善。第一，模型精度依赖更多真实田间样本，后续应补充不同季节、光照、树龄和病害阶段的图像；"
        "第二，大模型生成内容必须由教师审核，不能替代农技专家判断；第三，物联网环境数据接入还可进一步自动化，使温湿度、"
        "光照、土壤湿度等数据直接进入防治推理和教学案例；第四，应用效果还应继续量化，例如统计备课时间缩短比例、学生识别"
        "准确率变化、实训作品质量和课堂反馈。未来将建设案例库和学生作品库，推动该平台从单一红橙场景拓展到柑橘、茶叶、"
        "蔬菜等更多涉农实训场景。"
    )

    add_text(doc, f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}", "仿宋_GB2312", 16, center=False)
    doc.save(OUT)
    print(f"generated: {OUT}")


if __name__ == "__main__":
    generate()

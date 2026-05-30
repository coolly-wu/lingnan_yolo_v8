from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
PROJECT = "基于特色创新项目的岭南红橙病虫害智能监测与校本教研一体化平台"
TODAY = datetime.now().strftime("%Y年%m月%d日")


DISEASES = [
    ("柑橘黄龙病", "极高", "斑驳黄化、青果、红鼻子果；确诊后砍除并就地烧毁"),
    ("柑橘溃疡病", "高", "隆起火山口状褐色木栓化病斑"),
    ("柑橘疮痂病", "中", "黄褐色漏斗状或尖锥状疮痂"),
    ("柑橘炭疽病", "高", "叶尖叶缘枯死斑、果实下陷褐色圆斑"),
    ("柑橘黑斑病", "中", "红褐色中央凹陷黑色边缘斑"),
    ("柑橘红蜘蛛", "高", "叶背或果面密集红色微型虫体"),
    ("柑橘木虱", "极高", "成虫尾翘、若虫群落，是黄龙病传播媒介"),
    ("柑橘潜叶蛾", "中", "嫩叶银白色蜿蜒隧道和卷叶"),
    ("柑橘蚜虫", "中", "嫩梢叶背绿色或黑色虫体聚集"),
    ("柑橘介壳虫", "中", "盾形或棉絮状固着虫体"),
    ("柑橘花蕾蛆", "中", "灯笼状膨大花蕾及内部幼虫"),
    ("柑橘黑刺粉虱", "中", "叶背黑色若虫、白色蜡圈、煤烟病污染"),
]


MODULES = [
    ("app.py", "桌面应用启动入口，调用 lingnan.__main__.main()"),
    ("lingnan/config.py", "全局路径、模型候选、12类病虫害、物候期和阈值配置"),
    ("lingnan/core/inferencer.py", "ONNX / Ultralytics / Mock 多后端推理"),
    ("lingnan/core/annotator.py", "中文标签、彩色框和检测结果可视化"),
    ("lingnan/core/severity.py", "严重程度聚合与 Green / Amber / Red 分级"),
    ("lingnan/core/dataset_manager.py", "数据集扫描、划分、校验和 YOLO 数据配置管理"),
    ("lingnan/core/image_collector.py", "搜索关键词采集与原始样本管理"),
    ("lingnan/core/training_visualizer.py", "训练结果曲线、混淆矩阵和预测样例解析"),
    ("lingnan/core/teaching.py", "教学案例和学生农技实训指导意见生成核心"),
    ("lingnan/core/llm_client.py", "OpenAI 兼容本地/云端大模型调用客户端"),
    ("lingnan/data/knowledge_base.py", "SQLite 绿色防治知识库初始化与查询"),
    ("lingnan/data/log_manager.py", "检测历史台账和筛选查询"),
    ("lingnan/data/farmer_manager.py", "农户与果园档案管理"),
    ("lingnan/data/excel_exporter.py", "Excel 检测报表导出"),
    ("lingnan/data/pdf_exporter.py", "PDF 分析报告导出"),
    ("lingnan/data/markdown_exporter.py", "教学教研 Markdown / PDF 导出辅助"),
    ("lingnan/ui/main_window.py", "主窗口、标签页装配和页面间信号联动"),
    ("lingnan/ui/detection_page.py", "智能检测页面"),
    ("lingnan/ui/camera_page.py", "摄像头和视频文件检测页面"),
    ("lingnan/ui/dataset_page.py", "数据标注与数据集管理页面"),
    ("lingnan/ui/training_page.py", "模型训练页面"),
    ("lingnan/ui/training_result_page.py", "训练结果可视化页面"),
    ("lingnan/ui/teaching_page.py", "AI 教研生成页面"),
    ("lingnan/ui/settings_page.py", "字体、阈值、保存目录、LLM 配置等设置"),
]


def set_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(11)
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def shade(cell, color: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def title(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(15)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"编制日期：{TODAY}")
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)


def meta(doc: Document, name: str, kind: str) -> None:
    table(
        doc,
        ["项目", "内容"],
        [
            ("项目名称", PROJECT),
            ("文档名称", name),
            ("文档类型", kind),
            ("版本", "V1.0"),
            ("依据文件", "《技术规范文档》V3.0、《可行性研究报告》V3.0、README、当前代码结构"),
            ("适用对象", "项目实施人员、教师、农技人员、系统维护人员和验收人员"),
        ],
    )


def save(doc: Document, filename: str) -> None:
    set_style(doc)
    doc.save(ROOT / filename)


def install_manual() -> None:
    doc = Document()
    title(doc, "安装手册")
    meta(doc, "安装手册", "部署与安装说明")

    doc.add_heading("1. 安装目标", level=1)
    doc.add_paragraph(
        "本文档用于指导在 Windows 10/11、统信 UOS、麒麟 OS V10 等环境中安装、启动和验证本平台。"
        "系统以本地化、零网络依赖为基本目标，核心能力包括 YOLOv8 病虫害识别、绿色防治处方、"
        "检测台账、报表导出、模型训练工具链和校本教研内容生成。"
    )

    doc.add_heading("2. 运行环境要求", level=1)
    table(
        doc,
        ["类别", "最低要求", "推荐配置"],
        [
            ("操作系统", "Windows 10/11 64位", "Windows 11 64位；兼容统信 UOS、麒麟 OS V10"),
            ("Python", "Python 3.11.x", "Python 3.11.9"),
            ("CPU", "Intel i3-4xxx 或同级", "Intel i5-10xxx 或更高"),
            ("内存", "4 GB", "8 GB 及以上"),
            ("摄像头", "普通 UVC USB 摄像头", "USB 数字显微镜或高拍仪"),
            ("磁盘空间", "2 GB 可用空间", "10 GB 以上，便于保存数据集和训练结果"),
            ("网络", "运行时不依赖外网", "仅安装依赖、云端大模型或代码推送时需要网络"),
        ],
    )

    doc.add_heading("3. 项目目录说明", level=1)
    table(
        doc,
        ["目录/文件", "用途"],
        [
            ("app.py", "桌面应用启动入口"),
            ("lingnan/", "主程序包，包含核心算法、数据层和 GUI 页面"),
            ("models/", "放置 YOLOv8 / ONNX / INT8 模型权重，默认不提交 Git"),
            ("knowledge/", "SQLite 绿色防治知识库，启动时可自动生成"),
            ("runtime_data/", "运行期数据，包括检测日志、导出报告、训练结果和缓存"),
            ("docs/", "技术规范文档、可行性研究报告和导入模板"),
            ("tools/", "训练、导出、量化、基准测试和演示图生成脚本"),
            ("tests/", "pytest 单元测试"),
            ("build.bat / build.spec / installer.iss", "Windows 打包和安装向导配置"),
        ],
    )

    doc.add_heading("4. 开发环境安装步骤", level=1)
    nums(
        doc,
        [
            "进入项目根目录，例如 D:\\yolov8_gui。",
            "确认已安装 Python 3.11。",
            "创建并激活虚拟环境。",
            "安装 requirements.txt 中声明的依赖。",
            "将训练好的模型文件放入 models/ 目录，或使用内置 MockSimulator 进行界面联调。",
            "运行 app.py 启动应用。",
        ],
    )
    code(doc, "py -3.11 -m venv .venv\n.venv\\Scripts\\activate\npip install --upgrade pip\npip install -r requirements.txt\npython app.py")

    doc.add_heading("5. 模型文件部署", level=1)
    doc.add_paragraph("系统启动时会按优先级自动选择第一个存在的模型文件。推荐部署 INT8 量化 ONNX 模型。")
    table(
        doc,
        ["优先级", "文件位置", "说明"],
        [
            ("1", "models/yolov8s_xh_best_int8.onnx", "推荐部署版，体积小、速度快"),
            ("2", "models/yolov8s_xh_best.onnx", "FP32 ONNX 模型"),
            ("3", "models/yolov8s_xh_best.pt", "Ultralytics PyTorch 权重"),
            ("4", "models/yolov8s.pt 或项目根目录 yolov8s.pt", "演示或迁移调试使用"),
            ("5", "无模型文件", "自动进入 MockSimulator，便于 UI 联调"),
        ],
    )

    doc.add_heading("6. 启动方式", level=1)
    table(
        doc,
        ["方式", "命令/操作", "说明"],
        [
            ("Python 启动", "python app.py", "适合开发与调试"),
            ("模块启动", "python -m lingnan", "调用主包入口"),
            ("Windows 脚本", "双击 run.bat", "自动优先使用 .venv\\Scripts\\python.exe"),
            ("Linux/macOS 脚本", "./run.sh", "适合类 Unix 环境"),
        ],
    )

    doc.add_heading("7. 打包安装", level=1)
    doc.add_paragraph("Windows 可使用 PyInstaller 生成绿色版程序，并可配合 Inno Setup 生成安装向导。")
    code(doc, "pip install pyinstaller\nbuild.bat\n# 或\npyinstaller build.spec")
    bullets(
        doc,
        [
            "绿色版目录：dist/XHGan/XHGan.exe。",
            "安装向导：dist/XHGan_Setup_x64.exe，需本机安装 Inno Setup。",
            "打包前应确认模型文件、知识库生成逻辑和运行目录权限。",
        ],
    )

    doc.add_heading("8. 安装后验证", level=1)
    nums(
        doc,
        [
            "启动程序，确认主窗口正常显示。",
            "进入“设置”页，检查字体倍率、阈值、保存目录等配置可读写。",
            "进入“智能检测”页，导入红橙图片或演示图，确认可输出检测框、严重程度和防治处方。",
            "进入“历史台账”页，确认检测记录可筛选并导出 Excel / PDF。",
            "进入“教学生成”页，基于最近一次检测生成教学案例或实训指导意见。",
            "断开网络后再次运行，确认本地检测和知识库推荐不受影响。",
        ],
    )

    doc.add_heading("9. 常见问题处理", level=1)
    table(
        doc,
        ["问题", "可能原因", "处理建议"],
        [
            ("启动失败", "Python 版本或依赖缺失", "确认使用 Python 3.11，并重新安装 requirements.txt"),
            ("无检测结果", "未放置专属模型或阈值过高", "检查 models/ 目录；降低 conf 阈值；使用 MockSimulator 验证界面"),
            ("摄像头打不开", "设备被占用或驱动不兼容", "关闭其他摄像头软件，更换 UVC 摄像头或端口"),
            ("无法导出 PDF", "字体或 reportlab 依赖问题", "确认 reportlab 已安装，改用系统字体测试"),
            ("大模型生成不可用", "未配置 base_url/model 或网络不可达", "在设置页完善 LLM 配置；不可用时系统自动用模板生成"),
        ],
    )

    doc.add_heading("10. 安全与合规提示", level=1)
    bullets(
        doc,
        [
            "系统默认本地运行，图片、台账和导出报告均存放在 runtime_data/。",
            "模型权重、运行日志和检测数据库不建议提交到公共 Git 仓库。",
            "药剂推荐仅作教学和农技辅助，最终用药应服从当地农业农村部门最新指导。",
            "黄龙病确诊植株必须砍除并就地烧毁，禁止化学治疗。",
        ],
    )
    save(doc, "安装手册.docx")


def development_record() -> None:
    doc = Document()
    title(doc, "项目开发记录")
    meta(doc, "项目开发记录", "研发过程与工程实现记录")

    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph(
        "本项目面向岭南红橙产业园、合作社、农业技术推广和校本教研场景，构建“图像识别、绿色防治、"
        "台账追溯、教学生成”一体化平台。系统以 YOLOv8 本地推理为基础，结合 SQLite 知识库、"
        "PyQt5 桌面界面、报表导出和可选大模型内容生成能力，形成从田间图像到教学案例的完整闭环。"
    )

    doc.add_heading("2. 研发依据", level=1)
    bullets(
        doc,
        [
            "《技术规范文档》V3.0：按需求、数据、预处理、选型、训练、评估、优化、部署、监控九阶段组织。",
            "《可行性研究报告》V3.0：从技术、经济、操作、法律、进度五维评估，综合可行性评分 4.8 / 5。",
            "README：记录快速开始、模型加载、训练工具链、测试和打包流程。",
            "当前代码结构：主包已命名为 lingnan，新增教学教研扩展模块。",
        ],
    )

    doc.add_heading("3. 技术路线", level=1)
    table(
        doc,
        ["阶段", "主要工作", "工程产物"],
        [
            ("需求分析", "定义 12 类病虫害、5 个物候期、离线运行和适老化要求", "技术规范、KPI、FR/NFR 列表"),
            ("数据准备", "规划实地采集、公开图库辅助、YOLO 标注和数据划分", "数据集目录规范、classes.yaml/data.yaml"),
            ("模型选择", "选择 YOLOv8s-XH，强化小目标识别和 CPU 部署", "tools/yolov8s_xh.yaml"),
            ("训练评估", "训练、评估、误差分析和可视化", "tools/train.py、training_result_page.py"),
            ("模型优化", "ONNX 导出、INT8 量化、基准测试", "export_onnx.py、quantize_int8.py、benchmark.py"),
            ("桌面部署", "PyQt5 + Fluent 风格界面、本地知识库和台账", "lingnan/ui、lingnan/data"),
            ("教研扩展", "检测结果转教学案例/实训指导，支持 LLM 和模板降级", "teaching.py、llm_client.py、teaching_page.py"),
            ("测试交付", "pytest 覆盖核心逻辑和导出能力", "tests/"),
        ],
    )

    doc.add_heading("4. 里程碑记录", level=1)
    table(
        doc,
        ["里程碑", "完成内容"],
        [
            ("需求分析", "明确 12 类病虫害、五个物候期、离线运行、适老化界面、检测台账、教研生成等需求。"),
            ("数据准备", "规划岭南红橙田间图像、公开图库辅助样本、YOLO 标注格式、训练/验证/测试划分。"),
            ("模型选型", "选择 YOLOv8s-XH 作为核心检测模型，兼顾精度、小目标识别、CPU 部署与工程生态。"),
            ("核心推理", "实现模型候选自动加载、无模型时 MockSimulator 回退、中文路径读图和批量检测。"),
            ("知识库与台账", "建立 SQLite 防治处方、检测日志、农户档案、Excel/PDF 导出能力。"),
            ("图形界面", "完成 PyQt5 + Fluent 风格主界面，覆盖智能检测、摄像头、历史台账、设置、帮助等页面。"),
            ("数据与训练工具", "实现搜索采集、数据集管理、一键训练、ONNX 导出、INT8 量化、性能基准和训练结果可视化。"),
            ("教研扩展", "新增教学案例/学生农技实训指导生成，支持本地模板降级和 OpenAI 兼容大模型增强。"),
            ("测试与交付", "编写 pytest 测试，覆盖严重程度、知识库、日志、设置、推理、数据集、教研生成等模块。"),
        ],
    )

    doc.add_heading("5. 模块开发记录", level=1)
    table(doc, ["模块/文件", "职责说明"], MODULES)

    doc.add_heading("6. 核心功能实现", level=1)
    doc.add_heading("6.1 本地智能检测", level=2)
    bullets(
        doc,
        [
            "支持本地图片、文件夹、视频和 UVC 摄像头输入。",
            "按模型候选优先级自动加载 ONNX / PT 模型。",
            "无专属模型时自动进入 MockSimulator，保障界面和流程可演示。",
            "检测结果包含类别、置信度、边界框、严重程度、主病害和统计汇总。",
        ],
    )
    doc.add_heading("6.2 绿色防治与台账", level=2)
    bullets(
        doc,
        [
            "知识库以 12 类病虫害 × 5 个物候期生成约 60 组绿色防治处方。",
            "处方包含物理防治、生物防治、低毒化学药剂和 PHI 安全间隔期提示。",
            "检测日志写入 SQLite，可按时间、农户、病虫害类别筛选。",
            "支持 Excel 报表和 PDF 分析报告导出。",
        ],
    )
    doc.add_heading("6.3 校本教研生成", level=2)
    bullets(
        doc,
        [
            "检测结果与知识库处方被归一化为教学上下文。",
            "可生成《岭南红橙实时防害教学案例》和《学生农技实训指导意见》。",
            "默认使用模板生成，确保离线可用。",
            "可选接入 OpenAI 兼容本地/云端大模型，增强内容表达和课堂问题设计。",
            "教师审核后再导出，避免未审内容直接用于教学。",
        ],
    )

    doc.add_heading("7. 测试记录", level=1)
    table(
        doc,
        ["测试文件", "覆盖内容"],
        [
            ("tests/test_severity.py", "严重程度分级、黄龙病高危标记、面积/计数阈值"),
            ("tests/test_knowledge_base.py", "知识库初始化、60 组合查询和化学处方字段完整性"),
            ("tests/test_log_manager.py", "检测日志、农户档案 CRUD 和筛选"),
            ("tests/test_excel_exporter.py", "Excel 报表结构和导出结果"),
            ("tests/test_inferencer.py", "无模型时 Mock 后端回退"),
            ("tests/test_dataset_manager.py", "数据集扫描、划分、校验和配置生成"),
            ("tests/test_image_collector.py", "搜索采集和样本清单管理"),
            ("tests/test_training_visualizer.py", "训练曲线、混淆矩阵、预测样例解析"),
            ("tests/test_teaching.py", "教学上下文、Prompt、模板生成"),
            ("tests/test_markdown_exporter.py", "教学文档 Markdown 导出"),
        ],
    )

    doc.add_heading("8. 关键设计决策", level=1)
    table(
        doc,
        ["决策", "原因"],
        [
            ("选择 YOLOv8s-XH", "兼顾实时性、小目标识别、部署生态和 45 天工期可控性"),
            ("采用本地 SQLite", "降低部署门槛，满足离线运行和本地留痕"),
            ("采用 PyQt5 桌面客户端", "适合基层电脑部署、摄像头接入和离线使用"),
            ("保留 MockSimulator", "在无模型或演示环境下仍能展示完整业务流程"),
            ("教研生成采用模板优先", "保证零网络场景可用，大模型作为增强能力"),
            ("运行数据与模型权重不提交 Git", "避免仓库过大并保护本地数据资产"),
        ],
    )

    doc.add_heading("9. 当前状态与后续计划", level=1)
    bullets(
        doc,
        [
            "当前已形成桌面端核心功能、数据层、训练工具链、报表导出和教研生成模块。",
            "后续可增强物联网温湿度数据接入，将环境数据加入教学案例和防治策略。",
            "后续可接入 DeepSeek、豆包或本地模型服务，形成“识别 + 环境 + 大模型 + 教研”闭环。",
            "后续可完善安装包签名、自动更新、模型版本管理和现场反馈回流机制。",
        ],
    )
    save(doc, "项目开发记录.docx")


def user_manual() -> None:
    doc = Document()
    title(doc, "项目使用手册")
    meta(doc, "项目使用手册", "用户操作与业务流程说明")

    doc.add_heading("1. 系统简介", level=1)
    doc.add_paragraph(
        "本系统用于岭南红橙病虫害本地智能识别、绿色防治方案推送、检测台账管理、报表导出和校本教研内容生成。"
        "用户可通过图片、摄像头或视频文件完成检测，系统自动输出病虫害类别、置信度、严重程度、处方建议和教学材料。"
    )

    doc.add_heading("2. 主要功能", level=1)
    bullets(
        doc,
        [
            "12 类岭南红橙高发病虫害识别。",
            "红蜘蛛、蚜虫、木虱、黑刺粉虱等群聚虫害自动计数。",
            "轻度/中度/重度严重程度自动分级。",
            "按物候期联动三位一体绿色防治方案。",
            "黄龙病等高危病害突出预警。",
            "检测日志、农户/果园档案、Excel 和 PDF 导出。",
            "数据采集、数据标注、模型训练和训练结果可视化。",
            "教学案例和学生农技实训指导意见生成。",
        ],
    )

    doc.add_heading("3. 支持识别的病虫害", level=1)
    table(doc, ["病虫害", "危害等级", "主要识别特征/处理提示"], DISEASES)

    doc.add_heading("4. 快速上手流程", level=1)
    nums(
        doc,
        [
            "启动程序。",
            "在“农户档案”页录入农户、果园区块和联系方式，也可先跳过。",
            "在“智能检测”页选择当前物候期。",
            "导入红橙叶片、枝条、果面或花蕾图片。",
            "点击“开始检测”。",
            "查看右侧检测框、主病害、严重程度和绿色防治方案。",
            "需要留痕时保存台账，并在“历史台账”中导出 Excel 或 PDF。",
            "需要教学材料时进入“教学生成”页，生成并审核教学案例或实训指导意见。",
        ],
    )

    doc.add_heading("5. 智能检测页面", level=1)
    bullets(
        doc,
        [
            "选择物候期：春梢/萌芽、开花/谢花、生理落果/幼果、果实膨大/挂果、采收/冬剪。",
            "导入图片：支持 JPG、JPEG、PNG、BMP，兼容中文路径。",
            "调整阈值：conf 影响最低置信度，IoU 影响框合并；一般保持默认即可。",
            "查看结果：系统显示检测框、中文标签、置信度、主病害、严重程度和处方。",
            "台账保存：开启自动写台账后，检测结果会进入 runtime_data/detection_log.db。",
        ],
    )

    doc.add_heading("6. 摄像头与视频检测", level=1)
    bullets(
        doc,
        [
            "连接 UVC USB 摄像头、数字显微镜或高拍仪。",
            "在摄像头页面选择设备并启动预览。",
            "将病叶、果面或叶背虫害置于镜头下，保持光照稳定。",
            "系统实时显示检测结果，可用于课堂演示和田间快速筛查。",
            "导入视频文件后，可逐帧检测并导出带框视频。",
        ],
    )

    doc.add_heading("7. 历史台账与报表导出", level=1)
    bullets(
        doc,
        [
            "历史台账可按时间、农户、果园、病虫害类别筛选。",
            "Excel 报表适合日常统计和提交检查。",
            "PDF 分析报告适合形成阶段总结，包含统计、TOP 病害和防治建议。",
            "导出路径默认位于 runtime_data/exports，可在设置中修改。",
        ],
    )

    doc.add_heading("8. 农户与果园档案", level=1)
    bullets(
        doc,
        [
            "支持新增、编辑、删除农户和果园区块信息。",
            "检测时可选择农户档案，便于后续筛选和追溯。",
            "支持使用 CSV 模板导入基础档案。",
            "档案数据仅保存在本地。",
        ],
    )

    doc.add_heading("9. 数据采集、标注与训练", level=1)
    bullets(
        doc,
        [
            "数据采集页用于按搜索关键词采集原始样本，样本进入未分类目录。",
            "数据标注阶段再绑定 12 类病虫害类别，生成 YOLO 标签。",
            "训练页支持选择数据集、模型结构和训练参数，一键启动 YOLOv8 训练。",
            "训练结束后系统可发布 best.pt 到 models/ 目录。",
            "训练结果页自动解析 results.csv、results.png、混淆矩阵和预测示例图。",
        ],
    )

    doc.add_heading("10. 教学教研生成", level=1)
    doc.add_paragraph("教学生成页基于最近一次检测结果和绿色防治知识库，生成面向课堂和实训的教学材料。")
    nums(
        doc,
        [
            "先在智能检测页完成一次有效检测。",
            "进入“教学生成”页，确认检测素材摘要。",
            "选择文档类型：“岭南红橙实时防害教学案例”或“学生农技实训指导意见”。",
            "点击生成草稿。默认使用本地模板；若设置页已配置大模型，则可使用 LLM 智能生成。",
            "教师检查并修改草稿。",
            "勾选“我已审核，可导出”。",
            "导出 Markdown 或 PDF，用于课堂讲解、实训任务单或教研归档。",
        ],
    )
    table(
        doc,
        ["生成方式", "适用场景", "特点"],
        [
            ("模板生成", "断网、本地教学、稳定交付", "无需网络，结构固定，内容来自检测结果和知识库"),
            ("本地大模型", "校内私有化部署", "数据不出本机或局域网，表达更丰富"),
            ("云端 API", "有网络且已配置 DeepSeek/豆包等服务", "生成能力强，但会产生网络依赖和合规要求"),
        ],
    )

    doc.add_heading("11. 设置说明", level=1)
    table(
        doc,
        ["设置项", "说明"],
        [
            ("字体倍率", "调整界面字号，适合适老化展示"),
            ("默认 conf / IoU", "设置检测默认阈值"),
            ("默认物候期", "启动时自动选择常用物候期"),
            ("默认保存目录", "设置 Excel、PDF、教学文档等导出位置"),
            ("自动写台账", "开启后检测结果自动保存到本地数据库"),
            ("提示音", "重度病害或高危病害时提示"),
            ("LLM 开关", "启用或关闭大模型智能生成"),
            ("LLM base_url / model / api_key", "配置 OpenAI 兼容本地或云端大模型服务"),
        ],
    )

    doc.add_heading("12. 绿色防治注意事项", level=1)
    bullets(
        doc,
        [
            "优先采用农业防治、物理诱控、生物防治和生态调控措施。",
            "化学药剂必须关注安全间隔期 PHI，不得采前违规施药。",
            "黄龙病确诊植株必须立即砍除并就地烧毁，禁止化学治疗。",
            "本系统输出用于辅助判断和教学演示，最终处理应结合现场农技人员意见。",
        ],
    )

    doc.add_heading("13. 常见问题", level=1)
    table(
        doc,
        ["问题", "解决方法"],
        [
            ("导入图片后没有检测框", "检查模型是否加载、阈值是否过高、图片是否清晰。"),
            ("检测结果不准确", "换用更清晰的近景图；补充数据并重新训练专属模型。"),
            ("无法保存台账", "确认 runtime_data/ 目录有写入权限。"),
            ("教学生成提示无素材", "先在智能检测页完成一次检测。"),
            ("LLM 生成失败", "检查设置页中的 base_url、model、api_key 和网络；系统可自动降级为模板生成。"),
            ("PDF/Excel 找不到", "检查默认保存目录或 runtime_data/exports。"),
        ],
    )
    save(doc, "项目使用手册.docx")


def main() -> None:
    install_manual()
    development_record()
    user_manual()
    print("generated: 安装手册.docx, 项目开发记录.docx, 项目使用手册.docx")


if __name__ == "__main__":
    main()
